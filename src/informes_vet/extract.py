"""Parser de .docx -> dict estructurado.

Maneja:
- Tabla del paciente: 6×6 o 7×7 (con fila vacía inicial que se descarta).
- Merges horizontales: dedup de celdas consecutivas iguales.
- Extracción por etiqueta (regex) en vez de por índice de fila.
- 3 secciones por encabezado: hallazgos, conclusiones, firma.
- Fallback gestacional cuando no se detectan órganos canónicos.
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

import docx

from .hashutil import hash_doc, hash_hallazgo
from .organs import classify_state, segment_findings_smart

FIELD_LABELS: dict[str, str] = {
    "nombre": r"^(?:Nombre|Paciente)[:\s]*$",
    "especie": r"^Especie[:\s]*$",
    "raza": r"^Raza[:\s]*$",
    "genero": r"^(?:G[eé]nero|Sexo)[:\s]*$",
    "edad": r"^Edad[:\s]*$",
    "peso": r"^Peso[:\s]*$",
    "tutor": r"^Tutor[:\s]*$",
    "doctor_solicitante": r"^Dr\.?\s*(?:Solicitante)?[:\s]*$",
    "fecha": r"^Fecha[:\s]*$",
    "antecedentes": r"^Antecedentes[:\s]*$",
    "n_ficha": r"^N[°º\.]?\s*Ficha[:\s]*$",
    "motivo": r"^Motivo[:\s]*$",
    "anamnesis": r"^Anamnesis[:\s]*$",
    "estudio": r"^Estudio[:\s]*$",
}

_HEADER_FINDINGS = re.compile(
    r"^Descripci[oó]n\s+Hallazgos\s+Ecogr[aá]ficos:?\s*$", re.IGNORECASE
)
# Prioridad 1: header canónico estricto. Evita falsos positivos cuando el
# archivo tiene también "Conclusiones:" suelto más arriba.
_HEADER_CONCLUSIONS_STRICT = re.compile(
    r"^Conclusiones\s+Ecogr[aá]ficas:?\s*$", re.IGNORECASE
)
# Prioridad 2: header alternativo "Conclusiones:" / "CONCLUSIONES:" sin más.
# Solo se usa si el estricto no matchea en ningún párrafo.
_HEADER_CONCLUSIONS_LOOSE = re.compile(
    r"^Conclusiones\s*:?\s*$", re.IGNORECASE
)
# Prioridad 3: header con conclusión pegada (ej. "Conclusiones: Útero y
# ovarios...") o con prefijo no-alfanumérico (ej. "}Conclusiones..."). Usa
# \b para tolerar prefijos no-palabra.
_HEADER_CONCLUSIONS_TAIL = re.compile(
    r"\bConclusiones(?:\s+Ecogr[aá]ficas)?\s*:\s*(.*)$", re.IGNORECASE
)
_HEADER_SIG = re.compile(
    r"^(?:Atte\.?|Atentamente|MV\b|Dr\.?\s|M[eé]d[ií]co\sVet|"
    r"RUT\b|Socio\sSOCHIEPA|Diploma\sImagenolog)" r"[\s\-:\.]*",
    re.IGNORECASE,
)

_NULL_FIELDS = {"", None}

# Labels planos (sin regex) para detectar campos pegados por tab en valores
# capturados por _extract_from_paragraphs (caso Mía: "Nombre\tEspecie:").
_TAB_LABEL_NAMES = (
    "Nombre", "Paciente", "Especie", "Raza", "Genero", "Género", "Sexo",
    "Edad", "Peso", "Tutor", "Doctor", "Dr", "Dr. Solicitante", "Fecha",
    "Antecedentes", "N° Ficha", "N. Ficha", "Motivo", "Anamnesis", "Estudio",
)
_TAB_CONTINUATION = re.compile(
    r"\t\s*(?:" + "|".join(re.escape(n) for n in _TAB_LABEL_NAMES) + r")\s*[:=]",
    re.IGNORECASE,
)


def _clean_tab_continuation(value: str) -> str:
    """Recorta '\\tLabel:' o '\\s+Label:' pegado al final de un valor capturado.

    Caso Mía Particual: una línea trae 'Nombre: Mía\\tEspecie:\\tCanino'.
    Sin esta limpieza, 'nombre' queda con valor 'Mía\\tEspecie:\\tCanino'.
    """
    if "\t" not in value:
        return value
    m = _TAB_CONTINUATION.search(value)
    if m:
        return value[: m.start()].rstrip()
    return value


def _is_punctuation_only(value: str) -> bool:
    """True si el valor no tiene letras/dígitos (solo ':' o símbolos).

    Evita que la regex de doctor_solicitante capture 'Solicitante:' (sin valor)
    cuando el párrafo es 'Dr. Solicitante:' (caso Mía Particual).
    """
    return not any(ch.isalnum() for ch in value)


def _dedupe_row_cells(row_cells: list[str]) -> list[str]:
    """Elimina celdas consecutivas duplicadas (efecto de merges horizontales)."""
    out: list[str] = []
    for c in row_cells:
        c = (c or "").strip()
        if c and (not out or out[-1] != c):
            out.append(c)
    return out


def _is_row_empty(row_cells: list[str]) -> bool:
    return all(not (c or "").strip() for c in row_cells)


def _normalize_table(table: docx.table.Table) -> list[list[str]]:
    """Convierte la tabla a lista de filas con celdas deduplicadas."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells_raw = [cell.text for cell in row.cells]
        cells = _dedupe_row_cells(cells_raw)
        rows.append(cells)
    return rows


def _extract_from_table(rows: list[list[str]]) -> dict[str, str]:
    """Extrae campos del paciente recorriendo celdas y matcheando etiquetas.

    Soporta dos formatos de tabla:

    - **Moderno (B)**: label y valor en celdas adyacentes tras dedupe.
      Ej: ``['Fecha:', '01/04/2026']``.
    - **Legacy (A)**: una sola celda con ``"Label: Value"`` resultado de merges
      horizontales donde la celda original cubre múltiples columnas y
      ``_dedupe_row_cells`` la colapsa. Ej: ``['Fecha: 11 de septiembre de 2024']``.

    El fallback legacy sólo se activa para campos que el formato moderno no
    extrajo, preservando la prioridad de la extracción moderna.
    """
    fields: dict[str, str] = {}
    label_keys = list(FIELD_LABELS.keys())
    label_patterns = {k: re.compile(p, re.IGNORECASE) for k, p in FIELD_LABELS.items()}

    for row in rows:
        i = 0
        while i < len(row) - 1:
            label_text = (row[i] or "").strip()
            value_text = (row[i + 1] or "").strip()
            if not label_text or not value_text:
                i += 1
                continue
            for key in label_keys:
                if key in fields:
                    continue
                if label_patterns[key].match(label_text):
                    fields[key] = value_text
                    break
            i += 1

    for row in rows:
        for cell in row:
            t = (cell or "").strip()
            if not t or ":" not in t:
                continue
            label_part, _, value_part = t.partition(":")
            label_part = label_part.strip()
            value_part = value_part.strip()
            if not label_part or not value_part:
                continue
            for key in label_keys:
                if key in fields:
                    continue
                if label_patterns[key].match(label_part):
                    fields[key] = value_part
                    break
    return fields


def _extract_from_paragraphs(paragraphs_text: list[str], existing: dict[str, str]) -> dict[str, str]:
    """Busca campos que no se encontraron en la tabla, escaneando párrafos.

    Útil para los campos que están en párrafos sueltos (no en la tabla) en
    algunas variantes del documento.
    """
    label_keys = [k for k in FIELD_LABELS if k not in existing]
    out = dict(existing)
    for raw in paragraphs_text:
        text = (raw or "").strip()
        if not text:
            continue
        for key in label_keys:
            if key in out:
                continue
            pattern = re.compile(
                rf"^(?:{FIELD_LABELS[key].rstrip('$').rstrip(':')})[:\s]+(.+)$",
                re.IGNORECASE,
            )
            m = pattern.match(text)
            if m:
                val = _clean_tab_continuation(m.group(1).strip())
                if val and not _is_punctuation_only(val):
                    out[key] = val
                break

    # Segundo pass: capturar campos que comparten renglón separados por tabs
    # (caso Mía Particual: "Nombre: Mía\tEspecie:\tCanino"). Si un campo
    # todavía no está en `out`, lo buscamos dentro de los tabs.
    if any(k not in out for k in label_keys):
        for raw in paragraphs_text:
            text = (raw or "").strip()
            if "\t" not in text:
                continue
            parts = [p.strip() for p in text.split("\t") if p.strip()]
            for part in parts:
                for key in label_keys:
                    if key in out:
                        continue
                    pattern = re.compile(
                        rf"^(?:{FIELD_LABELS[key].rstrip('$').rstrip(':')})[:\s]+(.+)$",
                        re.IGNORECASE,
                    )
                    m = pattern.match(part)
                    if m:
                        val = m.group(1).strip()
                        if val and not _is_punctuation_only(val):
                            out[key] = val
                            break
    return out


def _find_section_paragraphs(
    doc: docx.document.Document,
) -> tuple[int | None, int | None, int | None, str]:
    """Devuelve (idx_hallazgos, idx_conclusiones, idx_firma, resto_header_conc).

    Estrategia del header de conclusiones (3 prioridades, en orden):
    1. Canónico 'Conclusiones Ecográficas:' (puro o con resto) — usa
       re.search con start<=2 para tolerar prefijo '}' o similar.
    2. Alternativo suelto 'Conclusiones:' / 'CONCLUSIONES:' — solo si (1)
       no encontró. Misma tolerancia de prefijo.

    `resto_header_conc` es el texto pegado al header en el mismo párrafo;
    vacío si el header estaba solo.
    """
    idx_findings: int | None = None
    idx_conclusions: int | None = None
    idx_sig: int | None = None
    resto_header_conc = ""

    paragraphs = [(i, (p.text or "").strip()) for i, p in enumerate(doc.paragraphs)]

    def _try_match_conclusion(t: str) -> tuple[int, str] | None:
        m = _HEADER_CONCLUSIONS_TAIL.search(t)
        if not m or m.start() > 2:
            return None
        # group(0) = "Conclusiones" o "Conclusiones Ecográficas:" o "Conclusiones: resto"
        es_canonico = "ecogr" in m.group(0).lower()
        return (1 if es_canonico else 2, m.group(1).strip())

    # Pasada 1: header canónico (con "Ecográficas")
    for i, t in paragraphs:
        if not t:
            continue
        if idx_findings is None:
            if _HEADER_FINDINGS.match(t):
                idx_findings = i
            continue
        if i <= idx_findings:
            continue
        if idx_conclusions is not None:
            if _HEADER_SIG.match(t):
                idx_sig = i
                break
            continue
        result = _try_match_conclusion(t)
        if result and result[0] == 1:
            idx_conclusions = i
            resto_header_conc = result[1]
            continue
        if t and _HEADER_SIG.match(t):
            idx_sig = i
            break

    # Pasada 2: header suelto (sin "Ecográficas")
    if idx_findings is not None and idx_conclusions is None:
        for i, t in paragraphs:
            if not t or i <= idx_findings:
                continue
            if _HEADER_SIG.match(t):
                idx_sig = i
                break
            result = _try_match_conclusion(t)
            if result and result[0] == 2:
                idx_conclusions = i
                resto_header_conc = result[1]
                continue

    # Buscar firma si todavía no la tenemos
    if idx_conclusions is not None and idx_sig is None:
        for i, t in paragraphs:
            if i <= idx_conclusions:
                continue
            if t and _HEADER_SIG.match(t):
                idx_sig = i
                break

    return idx_findings, idx_conclusions, idx_sig, resto_header_conc


def _extract_findings_text(doc: docx.document.Document, idx_findings: int, idx_conclusions: int | None) -> str:
    """Concatena los párrafos entre 'Descripción Hallazgos' y 'Conclusiones'."""
    end = idx_conclusions if idx_conclusions is not None else len(doc.paragraphs)
    chunks: list[str] = []
    for i in range(idx_findings + 1, end):
        t = (doc.paragraphs[i].text or "").strip()
        if t:
            chunks.append(t)
    return " ".join(chunks)


def _extract_conclusions_text(
    doc: docx.document.Document,
    idx_conclusions: int,
    idx_sig: int | None,
    resto_header: str = "",
) -> str:
    """Concatena los párrafos entre 'Conclusiones' y la firma. Sin fragmentar.

    Si el header de conclusiones traía texto pegado en el mismo párrafo
    (caso 'Conclusiones: Útero y ovarios...'), ese resto se prepende.
    """
    end = idx_sig if idx_sig is not None else len(doc.paragraphs)
    chunks: list[str] = []
    if resto_header:
        chunks.append(resto_header)
    for i in range(idx_conclusions + 1, end):
        t = (doc.paragraphs[i].text or "").strip()
        if t and not _HEADER_SIG.match(t):
            chunks.append(t)
    return " ".join(chunks).strip()


class ExtractionError(Exception):
    """Error al parsear un .docx. El caller debe logear y continuar."""


def parse_docx(path: Path) -> dict[str, Any]:
    """Lee un .docx y devuelve un dict listo para persistir.

    El dict contiene todos los campos del informe + listas de hallazgos.
    No abre la BD; la persistencia es responsabilidad del caller.
    """
    path = Path(path)
    try:
        doc = docx.Document(path)
    except Exception as e:
        raise ExtractionError(f"No se pudo abrir {path}: {e}") from e

    table_fields: dict[str, str] = {}
    if doc.tables:
        table = doc.tables[0]
        rows = _normalize_table(table)
        if rows and _is_row_empty(rows[0]) and len(rows) > 6:
            rows = rows[1:]
        table_fields = _extract_from_table(rows)

    paragraphs_text = [p.text for p in doc.paragraphs]
    fields = _extract_from_paragraphs(paragraphs_text, table_fields)

    idx_findings, idx_conclusions, idx_sig, resto_header_conc = _find_section_paragraphs(doc)

    hallazgos_crudos = ""
    if idx_findings is not None:
        hallazgos_crudos = _extract_findings_text(doc, idx_findings, idx_conclusions)

    hallazgos: list[dict] = []
    if hallazgos_crudos:
        # Pre-resolver estudio para pasarlo al segmentador inteligente y aplicar
        # la salvaguarda gestacional.
        estudio_para_segmentador = (fields.get("estudio") or "").strip()
        hallazgos = segment_findings_smart(
            hallazgos_crudos, estudio=estudio_para_segmentador
        )
        for h in hallazgos:
            h["hallazgo_hash"] = hash_hallazgo(h["organo"], h["descripcion"])

    # Prioridad 1: campo "Estudio" del DOCX (verdad de la plantilla).
    # Solo si el DOCX no lo trae, se usa heurística como fallback.
    estudio_docx = (fields.get("estudio") or "").strip()
    # Re-resolver estudio si todavía no está decidido, usando los hallazgos
    # que segment_findings_smart acaba de producir.
    if not estudio_docx:
        if not hallazgos or all(h["organo"] == "Gestación" for h in hallazgos):
            estudio = "Gestacional"
        else:
            estudio = "Abdominal"
    else:
        estudio = estudio_docx

    conclusiones_texto = ""
    if idx_conclusions is not None:
        conclusiones_texto = _extract_conclusions_text(
            doc, idx_conclusions, idx_sig, resto_header_conc
        )

    paciente_json = json.dumps(
        {k: fields.get(k) for k in FIELD_LABELS if k in fields},
        ensure_ascii=False,
    )

    record: dict[str, Any] = {
        "archivo": path.name,
        "ruta_relativa": path.as_posix(),
        "anio": _year_from_path(path),
        "sha256": hash_doc(doc),
        "nombre": fields.get("nombre"),
        "especie": fields.get("especie"),
        "raza": fields.get("raza"),
        "genero": fields.get("genero"),
        "edad": fields.get("edad"),
        "peso": fields.get("peso"),
        "tutor": fields.get("tutor"),
        "doctor_solicitante": fields.get("doctor_solicitante"),
        "fecha": fields.get("fecha"),
        "antecedentes": fields.get("antecedentes"),
        "motivo": fields.get("motivo"),
        "anamnesis": fields.get("anamnesis"),
        "n_ficha": fields.get("n_ficha"),
        "estudio": estudio,
        "hallazgos_crudos": hallazgos_crudos or None,
        "paciente_json": paciente_json,
        "hallazgos": hallazgos,
        "conclusiones_texto": conclusiones_texto,
    }
    return record


def _year_from_path(path: Path) -> int:
    m = re.search(r"Ecograf[ií]a\s+(\d{4})", str(path))
    return int(m.group(1)) if m else 0
