"""SHA-256 sobre el texto canónico de un .docx para idempotencia."""

import hashlib
import re

import docx


def canonical_text(doc: docx.document.Document) -> str:
    """Concatena paragraphs.text + cells de tables[0], normalizado.

    Se omiten medios embebidos (imágenes, video) porque no afectan al texto clínico.
    """
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    if doc.tables:
        for row in doc.tables[0].rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    text = " ".join(parts)
    text = re.sub(r"\s+", " ", text)
    return text


def hash_doc(doc: docx.document.Document) -> str:
    """SHA-256 hex del texto canónico."""
    return hashlib.sha256(canonical_text(doc).encode("utf-8")).hexdigest()


def hash_hallazgo(organo: str, descripcion: str) -> str:
    """SHA-256 hex de organo+descripcion normalizados (para `hallazgos.hallazgo_hash`)."""
    payload = (organo or "").strip().lower() + (descripcion or "").strip()
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()
